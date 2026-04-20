from docx import Document
from langchain_core.documents import Document as LangchainDocument
import os
import fitz  # PyMuPDF
import pypandoc
import re
from .config import MIN_CHUNK_SIZE, MAX_CHUNK_SIZE_BEFORE_SPLIT, TITLE_MERGE_THRESHOLD

import shutil

# 在文件顶部，检查pandoc是否安装了
if shutil.which("pandoc") is None:
    raise EnvironmentError(
        "未找到 pandoc 可执行文件！\n"
        "请安装 pandoc: https://pandoc.org/installing.html\n"
        "安装后重启 VSCode 使 PATH 生效。"
    )


def _format_table_to_markdown(table):
    """将docx表格对象转换为Markdown格式的文本"""
    markdown_table = ""
    # 表头
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    markdown_table += "| " + " | ".join(headers) + " |\n"
    # 分隔线
    markdown_table += "| " + " | ".join(["---"] * len(headers)) + " |\n"
    # 表格内容
    for row in table.rows[1:]:
        row_content = [cell.text.strip() for cell in row.cells]
        markdown_table += "| " + " | ".join(row_content) + " |\n"
    return markdown_table

def _load_docx_document(file_path: str) -> list[LangchainDocument]:
    """
    加载并解析 .docx 文件，通过 pandoc 转换为 Markdown 后，按一级标题进行语义分块。
    """
    try:
        # 1. 使用 pypandoc 将 docx 转换为 markdown 字符串
        md_content = pypandoc.convert_file(file_path, 'gfm', format='docx')
    except Exception as e:
        raise RuntimeError(f"使用 pandoc 转换 '{file_path}' 失败: {e}\n"
                         "请确保 pandoc 已正确安装并已添加到系统PATH中。")

    # 2. 使用正则表达式按 Markdown 一级标题 (^# ) 分割文本
    # 修改正则表达式，只匹配一级标题（单个 # 后跟空格）
    chunks = re.split(r'(^#\s[^#].*)', md_content, flags=re.MULTILINE)
    
    content_parts = []
    current_chunk = ""
    
    for chunk in chunks:
        if not chunk.strip():
            continue
        
        # 如果 chunk 是一个一级标题 (以单个 # 开头，且后面不是 #)
        if re.match(r'^#\s[^#]', chunk):
            # 如果当前块有内容，先保存
            if current_chunk.strip():
                content_parts.append(current_chunk.strip())
            # 开始一个新块，以这个标题开头
            current_chunk = chunk
        else:
            # 如果不是一级标题，就追加到当前块
            current_chunk += chunk

    # 添加最后一个累积的块
    if current_chunk.strip():
        content_parts.append(current_chunk.strip())
    
    # 后处理：合并过短的首个块（通常是孤立的文档标题）
    if len(content_parts) >= 2 and len(content_parts[0]) < TITLE_MERGE_THRESHOLD:
        # 将第一个短块与第二个块合并
        merged_first_chunk = content_parts[0] + "\n\n" + content_parts[1]
        content_parts = [merged_first_chunk] + content_parts[2:]
        
    source_metadata = {"source": os.path.basename(file_path)}
    return [LangchainDocument(page_content=part, metadata=source_metadata) for part in content_parts]

def _load_pdf_document(file_path: str) -> list[LangchainDocument]:
    """加载并解析 .pdf 文件，基于章节标题进行智能分块"""
    if not os.path.exists(file_path) or not file_path.endswith(".pdf"):
        raise ValueError("无效的文件路径或文件类型，需要 .pdf 文件。")
        
    doc = fitz.open(file_path)
    
    # 1. 提取所有页面的文本，合并为一个大字符串
    full_text = ""
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text = page.get_text("text")
        if text.strip():
            full_text += text + "\n\n"  # 页面间用双换行分隔
    
    if not full_text.strip():
        return []
    
    # 2. 重新设计的智能分块策略：更精确的分割逻辑
    chapter_pattern = r'第[一二三四五六七八九十\d]+章[：:].*'
    
    # 找到所有章节位置
    chapter_matches = list(re.finditer(chapter_pattern, full_text, flags=re.MULTILINE | re.IGNORECASE))
    
    raw_chunks = []
    
    if not chapter_matches:
        # 如果没有找到章节，按页面分块
        raw_chunks = [full_text.strip()]
    else:
        # 根据章节位置精确分割
        start_pos = 0
        
        for i, match in enumerate(chapter_matches):
            if i == 0:
                # 第一个章节之前的内容（如前言）
                if match.start() > 0:
                    pre_content = full_text[start_pos:match.start()].strip()
                    if len(pre_content) >= 50:
                        raw_chunks.append(pre_content)
                start_pos = match.start()
            else:
                # 提取前一个章节的完整内容
                chapter_content = full_text[start_pos:match.start()].strip()
                if len(chapter_content) >= 50:
                    raw_chunks.append(chapter_content)
                start_pos = match.start()
        
        # 添加最后一个章节
        final_content = full_text[start_pos:].strip()
        if len(final_content) >= 50:
            raw_chunks.append(final_content)
    
    # 第二步：应用大小限制进行优化
    content_parts = []
    
    # 小节模式，用于进一步拆分过大的章节
    subsection_pattern = r'^\d+\.\d+\s+[^\n]+'  # 如 "2.1 招聘与录用"
    
    for chunk in raw_chunks:
        chunk_size = len(chunk)
        
        # 如果块太大，尝试按小节拆分
        if chunk_size > MAX_CHUNK_SIZE_BEFORE_SPLIT:
            # 找到所有小节位置
            subsection_matches = list(re.finditer(subsection_pattern, chunk, flags=re.MULTILINE))
            
            if len(subsection_matches) >= 2:  # 至少有2个小节才进行拆分
                sub_start = 0
                current_sub_chunk = ""
                
                for j, sub_match in enumerate(subsection_matches):
                    if j == 0:
                        # 章节标题到第一个小节的内容
                        header_content = chunk[sub_start:sub_match.start()].strip()
                        if header_content:
                            current_sub_chunk = header_content
                        sub_start = sub_match.start()
                    else:
                        # 提取前一个小节的内容
                        subsection_content = chunk[sub_start:sub_match.start()].strip()
                        
                        # 如果当前块加上新内容会超出限制，先保存当前块
                        if current_sub_chunk and len(current_sub_chunk) + len(subsection_content) > MAX_CHUNK_SIZE_BEFORE_SPLIT:
                            content_parts.append(current_sub_chunk.strip())
                            current_sub_chunk = subsection_content
                        else:
                            # 合并到当前块
                            if current_sub_chunk:
                                current_sub_chunk += "\n\n" + subsection_content
                            else:
                                current_sub_chunk = subsection_content
                        sub_start = sub_match.start()
                
                # 添加最后一个小节
                final_subsection = chunk[sub_start:].strip()
                if current_sub_chunk and len(current_sub_chunk) + len(final_subsection) > MAX_CHUNK_SIZE_BEFORE_SPLIT:
                    content_parts.append(current_sub_chunk.strip())
                    content_parts.append(final_subsection)
                else:
                    if current_sub_chunk:
                        current_sub_chunk += "\n\n" + final_subsection
                    else:
                        current_sub_chunk = final_subsection
                    content_parts.append(current_sub_chunk.strip())
            else:
                # 没有足够的小节，直接使用整个块
                content_parts.append(chunk)
        
        # 如果块大小合适，直接使用
        elif chunk_size >= MIN_CHUNK_SIZE:
            content_parts.append(chunk)
        
        # 如果块太小，尝试与下一个块合并（这里先暂存，后面统一处理）
        else:
            content_parts.append(chunk)
    
    # 第三步：合并过小的块
    final_parts = []
    i = 0
    while i < len(content_parts):
        current_chunk = content_parts[i]
        
        # 如果当前块太小，尝试与下一个块合并
        if len(current_chunk) < MIN_CHUNK_SIZE and i + 1 < len(content_parts):
            next_chunk = content_parts[i + 1]
            merged_chunk = current_chunk + "\n\n" + next_chunk
            
            # 如果合并后不会过大，就合并
            if len(merged_chunk) <= MAX_CHUNK_SIZE_BEFORE_SPLIT:
                final_parts.append(merged_chunk)
                i += 2  # 跳过下一个块
            else:
                # 合并后会过大，保持原样
                final_parts.append(current_chunk)
                i += 1
        else:
            final_parts.append(current_chunk)
            i += 1
    
    content_parts = final_parts
    
    # 如果效果不好，退回到按页分块
    if len(content_parts) <= 1:
        print(f"PDF章节分块效果不佳，退回到按页分块模式")
        content_parts = []
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text("text")
            if text.strip():
                content_parts.append(text.strip())
            
    source_metadata = {"source": os.path.basename(file_path)}
    return [LangchainDocument(page_content=part, metadata=source_metadata) for part in content_parts]

def load_document(file_path: str) -> list[LangchainDocument]:
    """
    根据文件扩展名，加载并解析支持的文档类型。
    
    参数:
        file_path (str): 文件的完整路径。

    返回:
        list[LangchainDocument]: 包含文档内容的 Langchain 文档对象列表。
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件未找到: {file_path}")

    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == ".docx":
        return _load_docx_document(file_path)
    elif file_extension == ".pdf":
        return _load_pdf_document(file_path)
    else:
        raise ValueError(f"不支持的文件类型: '{file_extension}'。目前仅支持 .docx 和 .pdf。")

# --- 旧函数保留，以防万一，但逻辑已移至 _load_docx_document ---
def load_docx_documents(file_path: str):
    """
    加载并解析 .docx 文件，将文本和表格内容提取为 Langchain 文档对象。
    (此函数现在代理到新的实现)
    """
    if not os.path.exists(file_path) or not file_path.endswith(".docx"):
        raise ValueError("无效的文件路径或文件类型，需要 .docx 文件。")
        
    # 调用新的内部函数
    return _load_docx_document(file_path)

if __name__ == '__main__':
    # 这是一个示例用法，需要您提供一个测试用的docx文件
    # 比如在项目根目录下创建一个 data/source_documents 文件夹，并放入文件
    
    from config import DOCUMENT_DIR
    
    # 确保目录存在
    os.makedirs(DOCUMENT_DIR, exist_ok=True)
    
    # 创建一个虚拟的docx文件用于测试
    try:
        doc = Document()
        doc.add_paragraph("这是第一段。")
        doc.add_paragraph("这是第二段，会和第一段合并。")
        doc.add_paragraph("") # 空行
        doc.add_paragraph("这是新的一块段落。")
        table_data = [
            ['职级', '一线城市', '二线城市'],
            ['P1-P3', '500', '400'],
            ['P4-P5', '700', '600']
        ]
        table = doc.add_table(rows=1, cols=3)
        for i, header in enumerate(table_data[0]):
            table.cell(0, i).text = header
        for row_data in table_data[1:]:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
        
        doc.add_paragraph("表格后面的段落。")
        
        test_file_path = os.path.join(DOCUMENT_DIR, "test_document.docx")
        doc.save(test_file_path)
        
        print(f"创建测试文件于: {test_file_path}")
        
        # 加载并解析文档
        documents = load_document(test_file_path)
        
        print("\n--- 解析出的 Langchain 文档对象 (通过主函数 load_document) ---")
        for i, doc in enumerate(documents):
            print(f"--- Document {i+1} ---")
            print(f"Content:\n{doc.page_content}")
            print(f"Metadata: {doc.metadata}\n")
            
    except Exception as e:
        print(f"执行示例时出错: {e}")
